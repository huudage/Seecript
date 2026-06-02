"""一次性脚本:把 Seecript 仓库内所有 LLM 调用节点 dump 成 docx。

输出: docs/LLM-NODES.docx
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "LLM-NODES.docx"


# ----------------------- 节点数据 -----------------------------

NODES: list[dict] = [
    # ---------------------------------------------------------------------
    {
        "no": 1,
        "title": "视频画像 (Video Understanding)",
        "module": "样例拆解 / decompose_agent._video_understand",
        "file": "server/app/services/agent/decompose_agent.py",
        "model": "Doubao-Seed-2.0-lite (多模态)",
        "method": "complete_multimodal",
        "purpose": (
            "视频拆解 v2 的『理解先行』步骤。在做段落切分之前,先让 LLM 看完整片关键帧 "
            "(均匀采样 ≤12 张) 给出整支视频的语义画像 — 原型 / 叙事概览 / 建议段落数 / 基调。"
            "这是后面 _segment_with_roles 切段的输入,避免硬塞 hook/body/cta 这种死模板。"
        ),
        "input": [
            "system: _UNDERSTAND_SYSTEM (短视频内容分析师)",
            "user: 视频风格类型 + 总时长 + 声音情况 + 代表性镜头列表 (#index 起止 | 口播 | tags)",
            "images: ≤12 张关键帧缩略图 + 用户参考图 (可选,≤6 张)",
        ],
        "output": (
            "JSON: {\"archetype\": str(≤20 字), \"narrative_summary\": str(≤80 字), "
            "\"suggested_segments\": int(3-6), \"tone\": str(≤15 字)}"
        ),
        "constraints": [
            "archetype 截断到 40 字",
            "narrative_summary 截断到 200 字",
            "suggested_segments clamp 到 [3, 6]",
            "tone 截断到 30 字",
            "字段缺失则按 video_type 走 fallback (4 种 archetype 兜底)",
        ],
        "options": [
            "archetype: free-text (示例: 艺术展宣传 / 带货种草 / 城市 Vlog / 信息可视化解释)",
            "tone: free-text (示例: 冷静克制 / 高燃热血 / 诙谐自嘲 / 庄重正式)",
            "suggested_segments ∈ {3,4,5,6}",
        ],
        "system_prompt": (
            "你是短视频内容分析师。看一组按时间排序的关键帧(配可能为空的口播),"
            "请对整支视频做语义画像。\n"
            "返回 JSON:{"
            "\"archetype\": str(≤20字, 一句话定性这视频的原型;例:『艺术展宣传』『带货种草』『城市Vlog』『信息可视化解释』), "
            "\"narrative_summary\": str(≤80字, 一段话讲清整支视频在说什么、怎么说), "
            "\"suggested_segments\": int(3-6, 你建议把视频切成几个叙事段落), "
            "\"tone\": str(≤15字, 基调;例:『冷静克制』『高燃热血』『诙谐自嘲』『庄重正式』)"
            "}。\n"
            "注意:不要套用固定模板。视频拍什么样就说什么样 —— 艺术展就是艺术展,不要硬说『钩子→主体→引导』。"
        ),
        "user_prompt_template": (
            "视频风格类型(仅供参考):{video_type}({hint})\n"
            "总时长:{total_duration:.1f} 秒\n"
            "声音情况:{voice_hint}  # 『有口播』或『纯 BGM/环境音,无口播』\n\n"
            "代表性镜头(共 {n_total} 个,采样 {n_sampled} 个):\n"
            "#{idx} {start:.1f}-{end:.1f}s | {speech} | tags: {tags}\n"
            "...\n\n"
            "(可选) 附带 {n_ref} 张『用户参考画面』 —— 来自用户素材库,不属于样例视频,"
            "仅作为用户希望对齐的视觉气质提示。请仍以样例镜头为主体做分析。"
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 2,
        "title": "关键帧批量打标 (Frame Tagging)",
        "module": "样例拆解 / decompose_agent._attach_frame_tags",
        "file": "server/app/services/agent/decompose_agent.py",
        "model": "Doubao-Seed-2.0-lite (多模态)",
        "method": "complete_multimodal",
        "purpose": (
            "为每个样例镜头的缩略图打 3-5 个画面标签 + 推断字幕样式。"
            "tags 后续用于段落切分时给 LLM 当语义上下文,字幕样式则汇总为 PackagingProfile.subtitle_style。"
            "8 张一批走多模态接口,降低 token / 限流。"
        ),
        "input": [
            "system: _FRAME_TAG_SYSTEM (短视频画面打标助手)",
            "user: 『请为以下 N 张关键帧打标。frame_id 用 'f-001'.. 这种 0 填充三位的格式。』",
            "images: 一批 ≤8 张关键帧缩略图",
        ],
        "output": "JSON: {\"frame_tags\": [{\"frame_id\": str, \"tags\": [str], \"subtitle_style\": str}, ...]}",
        "constraints": [
            "每帧 3-5 个 tag",
            "subtitle_style 三选一: 大字加描边 / 小字白底 / 无字幕",
            "数组长度不足时用 {} 补齐 (避免 i 越界)",
            "失败整批降级为空 tag,流水线不挂",
        ],
        "options": [
            "subtitle_style ∈ {大字加描边, 小字白底, 无字幕}",
            "tags: free-text (封面风格 / 转场类型 / 物体场景 / 构图风格)",
        ],
        "system_prompt": (
            "你是短视频画面打标助手。输入是一组按时间排序的关键帧。"
            "请按封面风格 / 转场类型 / 字幕样式 / 物体场景 四个维度,"
            "为每帧给 3-5 个标签,并判定字幕样式(大字加描边 / 小字白底 / 无字幕)。"
            "返回 JSON:{\"frame_tags\": [{\"frame_id\": str, \"tags\": [str], \"subtitle_style\": str}]}"
        ),
        "user_prompt_template": (
            "请为以下 {batch_size} 张关键帧打标。frame_id 用 'f-001'.. 这种 0 填充三位的格式。"
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 3,
        "title": "Shot 角色+主题切段 (Shot Role Segmentation)",
        "module": "样例拆解 / decompose_agent._segment_with_roles",
        "file": "server/app/services/agent/decompose_agent.py",
        "model": "Doubao-Seed-2.0-lite (多模态)",
        "method": "complete_multimodal",
        "purpose": (
            "把每个镜头打 role + theme,再合并相邻同 role 镜头成段落。"
            "比『LLM 直接给段落起止时间』更稳:段落总时长一定 = 镜头总时长,绝不超出视频长度。"
            "v2 重构后取代旧 video_type 三选一 (hook/body/cta · opening/climax/closing · intro/build/drop/outro) 的 9 个固定 kind。"
        ),
        "input": [
            "system: _SHOT_ROLE_SYSTEM (短视频结构分析师)",
            "user: 视频原型 + 叙事概览 + 基调 + 镜头总数 + 总时长 + 镜头列表 (idx 起止 | 口播 | tags)",
            "images: 全部镜头缩略图",
        ],
        "output": "JSON: {\"shot_roles\": [{\"shot_index\": int, \"role\": str, \"theme\": str}]}",
        "constraints": [
            "role 必须 ∈ {opening, development, climax, closing}",
            "第一个镜头必须 opening",
            "最后一个镜头必须 closing",
            "中间不允许 opening / closing",
            "整片至多 1 个 climax",
            "theme 中文 ≤10 字 (后端再截断到 20 字)",
            "数组长度 = 镜头数,按 shot_index 升序",
            "缺漏 shot 按位置兜底 (首 opening / 尾 closing / 中间 development)",
        ],
        "options": [
            "role ∈ {opening, development, climax, closing}",
            "theme: free-text 中文短标签 (示例: 展品揭幕 / 艺术家自述 / 行动呼吁)",
        ],
        "system_prompt": (
            "你是短视频结构分析师。给定视频画像和按时间排序的镜头列表,"
            "为**每个镜头**标注它在叙事中的角色和主题。\n\n"
            "角色(role)只能是以下 4 种之一:\n"
            "- opening: 开场(吸引注意/奠定基调)\n"
            "- development: 发展铺陈(信息展开/内容主体)\n"
            "- climax: 高潮(情绪/视觉/冲突顶点)\n"
            "- closing: 收尾(余韵/引导/落版)\n\n"
            "硬约束:\n"
            "1. 第一个镜头必须是 opening\n"
            "2. 最后一个镜头必须是 closing\n"
            "3. 中间镜头不能是 opening 或 closing\n"
            "4. 整支视频最多 1 个镜头标 climax(也可以没有)\n"
            "5. 相邻同 role 镜头会被合并为一个段落 —— 所以最终段落数 ≤ 镜头数\n\n"
            "theme: 中文短标签(≤10 字),反映这个镜头真实在讲什么 ——"
            "不要照抄 role,要从画面/口播内容里提炼。\n\n"
            "返回 JSON:{\"shot_roles\": [{\"shot_index\": int, \"role\": str, \"theme\": str}]}\n"
            "数组长度必须等于镜头数,按 shot_index 升序排列。"
        ),
        "user_prompt_template": (
            "视频原型:{understanding.archetype}\n"
            "叙事概览:{understanding.narrative_summary}\n"
            "基调:{understanding.tone}\n"
            "镜头总数:{n_shots}\n"
            "总时长:{total:.1f} 秒\n"
            "{voice_hint}\n\n"
            "镜头列表(请为每一个镜头给出 role + theme):\n"
            "{idx}: {start:.1f}-{end:.1f}s | {speech} | tags: {tags}\n"
            "..."
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 4,
        "title": "段落结构改编 (Plan Adaptation)",
        "module": "Compose / plan_agent.adapt_structure",
        "file": "server/app/services/agent/plan_agent.py",
        "model": "Doubao-Seed-2.0-lite (多模态 / 纯文本)",
        "method": "complete_multimodal (有参考图时) / complete (纯文本)",
        "purpose": (
            "把 1-2 个样例视频的真实段落骨架按用户主题(brief)+视频目的(video_goal)+创作设置 "
            "(平台 / 调性 / CTA / 关键词 / 目标时长) 改编成新视频的段落结构。"
            "允许增/删/合并/重排;同时为每段产出 content_description,告诉创作者本段画面+口播该呈什么。"
        ),
        "input": [
            "system: _ADAPT_SYSTEM (短视频结构改编师)",
            "user: 样例画像 + 创作者输入 (brief, video_goal) + 创作设置 + 原样例段落池 (flat list)",
            "images: 用户参考图 / 参考视频抽帧 (可选)",
        ],
        "output": (
            "JSON: {\"adapted_sections\": [{\"role\": str, \"theme\": str, "
            "\"content_description\": str(30-100 字), \"duration_seconds\": number, "
            "\"source_section_indices\": [int]}]}"
        ),
        "constraints": [
            "总段数 3-7",
            "首段 role=opening、末段 role=closing、中间皆 development",
            "整片至多 1 个 climax",
            "每段 duration_seconds ∈ [2.0, 30.0]",
            "所有段时长之和接近目标总时长 (±20% 内不动,超出按比例缩放再 clamp)",
            "theme ≤8 字 (后端截到 20 字),content_description ≤300 字",
            "段数 <3 触发 fallback (1:1 拷贝样例);>7 截断保留首/末/最先 climax",
        ],
        "options": [
            "role ∈ {opening, development, climax, closing}",
            "settings.target_platform ∈ {douyin(抖音 9:16), wechat(视频号 9:16), xiaohongshu(小红书), bilibili(B 站 16:9)}",
            "settings.tone ∈ {tight_hype(紧凑高燃), calm_narrative(沉稳叙事), casual_daily(轻松日常), professional_cool(专业冷静)}",
        ],
        "system_prompt": (
            "你是短视频结构改编师。给定 1-2 个参考样例视频的真实段落结构、视频画像,以及"
            "创作者的主题、视频目的与创作设置,请把这些样例的『骨架』改编为本次新视频的段落结构。\n\n"
            "若给了 2 个参考样例,请将它们作为对等的灵感来源,不必偏向某一份;可借用任一份的"
            "节奏、卡点与段落创意,但不要把两份段落简单拼接(最终段数仍受 3-7 段硬约束)。\n\n"
            "允许:增加段落、删除冗余段落、合并相邻段落、调整顺序。\n\n"
            "硬约束:\n"
            "1. 第一段 role 必须是 opening\n"
            "2. 最后一段 role 必须是 closing\n"
            "3. 整支视频最多 1 段 climax(可以没有)\n"
            "4. 中间段都是 development(不允许中间出现 opening/closing)\n"
            "5. 总段数 3-7\n"
            "6. 所有段 duration_seconds 之和必须接近『目标总时长』(±20% 以内)\n\n"
            "每段返回字段:\n"
            "- role: opening | development | climax | closing\n"
            "- theme: 中文短标签(≤8 字),紧贴创作者主题,不照抄样例\n"
            "- content_description: 内容说明(30-100 字) —— 告诉创作者画面该呈现什么、"
            "口播该说什么、为什么放在这个位置,紧扣 brief + video_goal;若给了关键词,"
            "尽量自然融入;若给了 CTA,closing 段口播须体现\n"
            "- duration_seconds: 本段时长(浮点秒)。opening/closing 各 3-5s,climax 5-10s,"
            "development 4-8s;所有段之和贴近目标总时长\n"
            "- source_section_indices: 改编自原样例池哪些段落下标(合并后的 flat 下标);纯新增段为 []\n\n"
            "返回 JSON:{\"adapted_sections\": [{\"role\": str, \"theme\": str, "
            "\"content_description\": str, \"duration_seconds\": number, "
            "\"source_section_indices\": [int]}]}"
        ),
        "user_prompt_template": (
            "样例视频画像:\n{understanding_text}\n\n"
            "创作者输入:\n"
            "- 主题/卖点(brief):{brief_text}\n"
            "- 视频要求与目的(video_goal):{goal_text}\n\n"
            "创作设置:\n"
            "- 目标总时长:{target_total:.0f}s\n"
            "- 目标平台:{platform_label}\n"
            "- 整体调性:{tone_label}\n"
            "- 核心 CTA:{cta_text}\n"
            "- 必须出现的关键词:{kw_text}\n\n"
            "原样例共 {n_src} 段:\n"
            "[{idx}] (样例A/B 标签) role={role} | theme={theme} | shots={shot_indices} | summary={summary}\n"
            "...\n\n"
            "请基于以上信息改编段落结构(3-7 段,遵守硬约束,所有段时长之和贴近 {target_total:.0f}s)。\n"
            "(可选) 附带 {n_ref} 张『参考画面』 —— 不是样例镜头,而是用户希望对齐的视觉风格/构图/调性参考。"
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 5,
        "title": "新素材打标 (Material Tagging)",
        "module": "Compose / routers/material._tag_with_llm",
        "file": "server/app/routers/material.py",
        "model": "Doubao-Seed-2.0-lite (多模态)",
        "method": "complete_multimodal",
        "purpose": (
            "用户上传新素材(image / video 抽首帧)后,LLM 看一帧给标签 + 推荐段落角色 + 高光评分。"
            "高光评分 (0-1) 后续在 gap_agent.detect_gaps 决定『高 impact 段先吃高光素材』的排序。"
        ),
        "input": [
            "system: _MATERIAL_TAG_SYSTEM (短视频素材打标 Agent)",
            "user: video_type + allowed_sections + media_type + 提示按 schema 返回 JSON",
            "images: [thumbnail_path] (image 用原图;video 用 ffmpeg t=0.5s 抽首帧)",
        ],
        "output": (
            "JSON: {\"tags\": [string](3-5 个), \"recommended_section\": string, "
            "\"highlight_score\": number(0.0-1.0), \"highlight_reason\": string(≤20 字)}"
        ),
        "constraints": [
            "tags 取前 5 个,每个截断到 30 字",
            "recommended_section 不在白名单时回落 development",
            "highlight_score 容错为 float,clamp 到 [0,1];缺失给 0.5",
            "highlight_reason 截断到 60 字;缺失给『LLM 未给理由』",
            "audio 类型直接走 placeholder,不调 LLM",
            "兼容 mock 路径返回 {frame_tags: [{...}]} 包裹结构",
        ],
        "options": [
            "recommended_section ∈ {opening, development, climax, closing}",
            "highlight_score 语义:0.8+ 强冲击 (开头/高潮);0.5-0.8 标准镜头 (中段);<0.5 仅 B-roll",
        ],
        "system_prompt": (
            "你是短视频素材打标 Agent。看一帧画面,返回 JSON:\n"
            "{\"tags\": [string](3-5 个,物体/场景/构图/风格关键词),"
            "\"recommended_section\": string(必须从 allowed_sections 里选一个 role:"
            "opening 适合做开场钩子/标题铺垫;development 适合主体铺陈/对比/演示;"
            "climax 适合视觉/情绪顶点强构图;closing 适合行动引导/落版/余韵),"
            "\"highlight_score\": number(0.0-1.0;0.8+ 强冲击/可做开头或高潮,"
            "0.5-0.8 标准镜头适合中段,<0.5 仅 B-roll),"
            "\"highlight_reason\": string(一句话理由:构图/动作/情绪/光线,≤20 字)}。\n"
            "字段名 frame_tags / material_tag 是 mock 路由用,不要漏。"
        ),
        "user_prompt_template": (
            "video_type={video_type}\n"
            "allowed_sections=['opening', 'development', 'climax', 'closing']\n"
            "media_type={media_type}\n"
            "请按 system 中的 schema 返回 JSON,highlight_score 必须给一个 0.0-1.0 的数。"
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 6,
        "title": "缺口文案补全 (Gap Copy Fill)",
        "module": "Compose / gap_agent.fill_gap (action='copy')",
        "file": "server/app/services/agent/gap_agent.py",
        "model": "Doubao-Seed-2.0-lite",
        "method": "complete_json",
        "purpose": (
            "缺口补全的三种动作之一(rerank / copy / aigc)。当用户对某个 Gap 选『采纳文案』时,"
            "LLM 基于视频整体背景 + 该段 content_description + 主题 + 创作者额外提示生成口播旁白。"
            "文案兼做大字 text_card 兜底,所以要求口播友好、字数严控、紧扣本段。"
        ),
        "input": [
            "system: _COPY_SYSTEM (短视频口播 + 文案作者)",
            "user: 视频整体背景(brief) + 视频目的(goal) + 本段内容要求 + 主题词 + 原始槽位需求 + 标签提示 + 创作者补充",
        ],
        "output": "JSON: {\"gap_fill_narration\": str, \"alternatives\": [str, str]}",
        "constraints": [
            "主文案 ≤40 字中文",
            "备选 2 句各 ≤40 字 (后端 alternatives 最多取 3)",
            "禁止出现段落角色名 (opening/development/climax/closing)",
            "禁止『本段』『第 X 段』等元数据词",
            "禁止 markdown / ASCII 引号",
            "失败兜底:『[fallback] 这里加一句口播,把刚才的对比强调一下。』",
        ],
        "options": [
            "无固定枚举;输出为自由口播文案",
            "锚点优先级:本段 content_description > 视频整体背景 brief > 创作者补充 prompt_hint",
        ],
        "system_prompt": (
            "你是短视频口播 + 文案作者。你的输出会被前端用作两种用途:\n"
            "  1) 段落口播旁白(≤40 字)\n"
            "  2) 当该段没有合适视频素材时,作为『大字画面 text_card』兜底 —— 把这句文案做成全屏字卡。\n"
            "\n"
            "因此文案必须『紧扣本段的内容要求』+『与整体视频背景一致』+『口播友好且适合做大字』。\n"
            "\n"
            "你会收到三类锚点信息(按优先级):\n"
            "  - 视频整体背景(brief):决定语气和品类\n"
            "  - 本段的内容要求(content_description):决定本句到底要讲什么 —— 最高优先\n"
            "  - 创作者补充(prompt_hint):用户在面板里手填的特殊要求,权重低于上面两条\n"
            "\n"
            "硬约束:\n"
            "  - 主文案 ≤ 40 字中文,备选 2 句各 ≤ 40 字\n"
            "  - 紧扣『内容要求』,不要泛化到整体背景层\n"
            "  - 不出现段落角色名(opening/development/climax/closing)和『本段』『第 X 段』等元数据词\n"
            "  - 不要 markdown / ASCII 引号\n"
            "\n"
            "返回 JSON:{\"gap_fill_narration\": str, \"alternatives\": [str, str]}。"
        ),
        "user_prompt_template": (
            "视频整体背景:{brief}\n"
            "视频目的:{goal}\n"
            "本段内容要求:{content_desc}\n"
            "本段主题词:{theme}\n"
            "原始槽位需求(兜底):{gap.requirement}\n"
            "可参考素材标签:{tag_hint}\n"
            "创作者补充(低优):{prompt_hint}\n"
            "请输出主文案 + 2 句备选,紧扣『本段内容要求』。"
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 7,
        "title": "AIGC T2V Prompt 生成 (Seedance Prompt Engineering)",
        "module": "Compose / aigc_prompt_agent.generate_aigc_prompt",
        "file": "server/app/services/agent/aigc_prompt_agent.py",
        "model": "Doubao-Seed-2.0-lite",
        "method": "complete (取回纯文本再 _extract_json)",
        "purpose": (
            "把段落上下文转写为 Seedance T2V 友好的完备 prompt。"
            "原 gap.requirement 是给创作者看的中文段落描述,缺少镜头/景别/光线/质感等 T2V 要素 → 出片偏差大。"
            "本 agent 输出的 prompt 直接送给 Seedance T2V client.submit() 生成 5-12s 短片。"
        ),
        "input": [
            "system: _PROMPT_SYSTEM (Seedance T2V 提示词工程师)",
            "user: 段落角色 + 主题 + 时长 + 内容说明 + 原始槽位需求 + 视频整体主题 + 视频目的 + 创作者额外提示",
        ],
        "output": "JSON: {\"prompt\": \"...一句完备的 t2v_prompt...\"}",
        "constraints": [
            "总长 60-120 字中文 (后端 _sanitize 截到 200 字)",
            "必须覆盖 6 要素:主体 / 景别 / 机位运动 / 光线与色调 / 质感 / 情绪氛围",
            "禁止出现段落角色名 (opening/development/climax/closing)",
            "禁止『本段』『第 X 段』『片段』等元数据词",
            "禁止 ASCII 引号 / markdown",
            "禁止把『时长 Ns』直接写进 prompt 文案",
            "失败兜底拼装 content_description + requirement + hint + 默认拍摄要素",
        ],
        "options": [
            "景别 ∈ {特写, 中景, 远景, 航拍, 大全景, ...}",
            "机位运动 ∈ {固定, 推进, 拉远, 跟随, 摇移, 手持}",
            "光线与色调 ∈ {黄昏暖光, 冷调高对比, 自然光, 棚拍硬光, ...}",
            "质感 ∈ {电影感, 纪实, 产品级, 杂志感, ...}",
            "情绪氛围 ∈ {紧张, 庄重, 轻快, 神秘, ...}",
        ],
        "system_prompt": (
            "你是 Seedance 文生视频(T2V)的提示词工程师。给定一个短视频段落的角色、主题、"
            "内容说明、时长,以及视频的整体主题与目的,请输出一句**完备的中文 t2v_prompt** ——"
            "Seedance 直接拿这一句去生成画面。\n\n"
            "要素必须覆盖(缺一不可):\n"
            "1. 主体:画面里的人/物,正在做什么\n"
            "2. 景别:特写 / 中景 / 远景 / 航拍 / 大全景,至少一个\n"
            "3. 机位运动:固定 / 推进 / 拉远 / 跟随 / 摇移 / 手持,选一个最合本段叙事的\n"
            "4. 光线与色调:黄昏暖光 / 冷调高对比 / 自然光 / 棚拍硬光 等\n"
            "5. 质感:电影感 / 纪实 / 产品级 / 杂志感 等\n"
            "6. 情绪/氛围:紧张 / 庄重 / 轻快 / 神秘 等\n\n"
            "硬约束:\n"
            "- 总长 60-120 字中文,一句话或两个短句\n"
            "- 不出现段落角色名(opening/development/climax/closing)\n"
            "- 不出现『本段』『第 X 段』『片段』等元数据词\n"
            "- 不要 ASCII 引号、不要 markdown\n"
            "- 不要把『时长 Ns』直接写进 prompt 文案(duration_seconds 由后端单独传给 Seedance)\n\n"
            "返回 JSON:{\"prompt\": \"...一句完备的 t2v_prompt...\"}"
        ),
        "user_prompt_template": (
            "段落角色:{role}\n"
            "段落主题:{theme}\n"
            "段落时长:约 {duration:.1f}s\n"
            "段落内容说明:{content_desc}\n"
            "原始槽位需求:{gap.requirement}\n"
            "视频整体主题:{brief}\n"
            "视频要求与目的:{goal}\n"
            "创作者额外提示:{hint}\n"
            "请输出一句完备的 t2v_prompt,覆盖主体/景别/机位/光线/质感/情绪。"
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 8,
        "title": "包装推荐:转场 + 封面 (Packaging Recommend)",
        "module": "Compose / packaging_agent.recommend_packaging",
        "file": "server/app/services/agent/packaging_agent.py",
        "model": "Doubao-Seed-2.0-lite",
        "method": "complete_json",
        "purpose": (
            "看完 plan.main_track (每段 role+theme+口播) 一次性给出:"
            "(a) 相邻段落切换处的转场风格;"
            "(b) 一份开场封面方案 (主标题/副标题/调色板/排版/风格说明)。"
            "system prompt 根据 PackagingPreferences 动态拼装:把 allowed_styles/max_duration/cover 策略 inline 进去,"
            "确保 LLM 输出在白名单内。"
        ),
        "input": [
            "system: _build_system_prompt(prefs) — 根据 prefs 动态拼装",
            "user: 创作者主题(brief) + video_goal + plan_id + 总时长 + 主轨分镜列表",
            "temperature: prefs.llm_temperature (来自 PackagingPreferences)",
        ],
        "output": (
            "JSON: {\"transitions\": [{\"at_seconds\": number, \"from_section\": str, \"to_section\": str, "
            "\"style\": str, \"duration\": number, \"reason\": str(≤30 字)}], "
            "\"cover\": {\"title\": str(≤12 字), \"subtitle\": str(≤18 字), "
            "\"palette\": [hex 2-3 个], \"layout\": str, \"style_note\": str(≤30 字)}}"
        ),
        "constraints": [
            "transition.style 必须 ∈ prefs.allowed_transition_styles (不在白名单替换为白名单首项)",
            "transition.duration ∈ [0.1, prefs.max_transition_duration] (clamp)",
            "transition.from_section / to_section 必须 ∈ {opening, development, climax, closing}",
            "transition.at_seconds 与 plan 真实段落切换点对齐 (±0.5s 容差;>0.5s 强制对齐)",
            "cover.title 按 cover_text_source 路由:custom 用 prefs.cover_custom_text;video_goal 用 plan.video_goal[:12];auto 用 LLM 给的 title",
            "cover.layout ∈ {center, left, split, stacked} (无效回落 center)",
            "cover.palette 必须是 hex(#RRGGBB),无效回落 [#FFE600, #1F2937]",
            "subtitle 长度 ≤18 字 (prefs.cover_with_subtitle=False 时强制 None)",
            "失败时规则兜底 (_RULE_TRANSITION 12 组 role 对 → 转场风格表 + 通用封面)",
        ],
        "options": [
            "transition.style ∈ {hard_cut, dissolve, slide, zoom, whip, wipe}",
            "from/to_section ∈ {opening, development, climax, closing}",
            "cover.layout ∈ {center, left, split, stacked}",
            "PackagingPreset ∈ {custom, minimalist, energetic, info_feed, dialogue}",
            "preset → allowed_transition_styles / max_duration / 字幕/封面 字段全部展开 (见 _PRESET_EXPANSIONS)",
            "cover_text_source ∈ {auto(LLM 自由发挥), video_goal(贴 plan.video_goal), custom(用户自定义)}",
        ],
        "system_prompt": (
            "[动态拼装,以下为模板;allowed/max_dur/cover_hint/bilingual_hint 在运行时填入]\n\n"
            "你是短视频包装设计师。根据给定的主轨分镜(每段标了 role+theme)与创作者主题文本,"
            "请输出两类建议:(a) 相邻段落切换处的转场风格;(b) 一份开场封面方案。\n"
            "转场只能从这些风格里选:[{allowed}],duration 必须 ≤ {max_dur:.2f}s。\n"
            "{cover_hint}。{bilingual_hint}\n"
            "返回 JSON:{"
            "\"transitions\": [{\"at_seconds\": number, \"from_section\": str, \"to_section\": str, "
            "\"style\": one of allowed, "
            "\"duration\": number (0.1-{max_dur:.2f}), "
            "\"reason\": str (≤30 字)}], "
            "\"cover\": {\"title\": str (≤12 字, 强冲击), \"subtitle\": str (≤18 字, 可空), "
            "\"palette\": [hex 颜色 2-3 个, 主色 + 强调色], "
            "\"layout\": one of [center, left, split, stacked], "
            "\"style_note\": str (≤30 字, 字号/色/排版)}"
            "}。\n"
            "from_section/to_section 必须是这 4 个 role 之一:opening / development / climax / closing。\n"
            "转场风格指导:opening→development 切到主体用节奏感强的;"
            "development→climax 进入高潮用冲击感强的;"
            "climax→closing 或 development→closing 切到收尾用情绪缓冲的。\n\n"
            "[cover_hint 三选一]\n"
            "  auto: 封面主标题由你自由发挥,强冲击 ≤12 字\n"
            "  video_goal: 封面主标题应紧贴用户的 video_goal 文本(你看到的『创作者主题』),≤12 字\n"
            "  custom: 封面主标题字段会被用户自定义文本替代,你给的 title 可被忽略;仍按 ≤12 字给一个候选\n\n"
            "[bilingual_hint 仅在 prefs.subtitle_bilingual=True 时拼入]\n"
            "  注意:本次开启双语字幕,封面 subtitle 字段请给一句英文翻译(≤20 字)。"
        ),
        "user_prompt_template": (
            "创作者主题:{plan.brief or '(创作者未提供主题文本)'}\n"
            "video_goal:{plan.video_goal or '(创作者未提供 video_goal)'}\n"
            "plan_id:{plan.plan_id}\n"
            "总时长:{plan.duration_seconds:.1f} 秒\n"
            "主轨分镜([role] 起止 · 口播):\n"
            "  - [{section}] {start:.1f}-{end:.1f}s · {narration or '(无口播)'}\n"
            "  - ..."
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 9,
        "title": "自然语言编辑 - 内容轨 (Edit · main track)",
        "module": "Render / routers/edit.apply_edit (track='main')",
        "file": "server/app/routers/edit.py",
        "model": "Doubao-Seed-2.0-lite",
        "method": "complete_with_tools (OpenAI tool calling)",
        "purpose": (
            "用户在底部 textarea 输入自然语言指令(如『把开场改 3 秒』『sc-1 加 dissolve 转场』),"
            "LLM 把意图翻译成原子 tool call,后端按顺序应用到 Plan,生成新 plan_id。"
            "三轨分流:main 只能改时长/素材/转场;packaging 只能改字幕/BGM 音量;voice 只能改口播 narration。"
            "渲染态锁:project.current_step=='render' 时 track=='main' 直接 409。"
        ),
        "input": [
            "system: _SYSTEM_MAIN (内容轨编辑助手,关键词→tool 路由表)",
            "user: 当前 main_track 列表(含 scene_id, section, src, dur, narr, transition_in) + 用户选区 marks + 用户指令",
            "tools: _TOOLS_MAIN (3 个原子 tool)",
        ],
        "output": (
            "OpenAI tool_calls: [{\"name\": \"edit_scene_duration\"|\"replace_scene_material\"|"
            "\"set_scene_transition\", \"arguments\": {...}}, ...]"
        ),
        "constraints": [
            "edit_scene_duration: duration 下限 0.5 秒",
            "set_scene_transition: style ∈ {hard_cut, dissolve, slide, zoom, whip, wipe},无效 → dissolve",
            "set_scene_transition: duration ∈ [0.1, 1.5],默认 0.4",
            "sc-0 (首段) 的 transition_in 调用会被忽略",
            "无任何 tool_call 命中 → 409,提示用户用更明确的指令",
            "渲染态锁:project.current_step=='render' → 409『内容轨不可改』",
        ],
        "options": [
            "tools 名 ∈ {edit_scene_duration, replace_scene_material, set_scene_transition}",
            "transition style ∈ {hard_cut, dissolve, slide, zoom, whip, wipe}",
            "关键词路由:『时长/更长/更短/缩短/拉长/N秒』→ edit_scene_duration;"
            "『替换/换成/改成/用素材』→ replace_scene_material;"
            "『转场/过渡/切换/dissolve/渐变/推拉/缩放/擦除』→ set_scene_transition",
        ],
        "system_prompt": (
            "你是视频剪辑助手,本次只能修改【内容轨】(main_track)。"
            "可选 tool:调整 scene 时长 / 替换 scene 素材 / 设置 scene 入场转场。"
            "禁止改字幕、BGM、口播 —— 那些是其他轨道的工具。"
            "『时长 / 更长 / 更短 / 缩短 / 拉长 / N秒』→ edit_scene_duration;"
            "『替换 / 换成 / 改成 / 用素材』→ replace_scene_material;"
            "『转场 / 过渡 / 切换 / dissolve / 渐变 / 推拉 / 缩放 / 擦除』→ set_scene_transition;"
            "set_scene_transition 的 style 必须取自 {hard_cut, dissolve, slide, zoom, whip, wipe},"
            "其他词都先归到 dissolve;duration 不填默认 0.4 秒(范围 0.1-1.5)。"
        ),
        "user_prompt_template": (
            "当前 Plan main_track:\n"
            "- {scene_id} ({section}) src={source_ref} dur={duration:.1f}s narr={narration!r} trans={style}/{dur:.2f}s\n"
            "- ...\n"
            "(可选) 用户选中:\n[{track}] {start:.1f}-{end:.1f}s target={target_id}\n"
            "用户指令(轨道=main):{instruction}\n"
            "输出 edit_tool_calls。"
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 10,
        "title": "自然语言编辑 - 包装轨 (Edit · packaging track)",
        "module": "Render / routers/edit.apply_edit (track='packaging')",
        "file": "server/app/routers/edit.py",
        "model": "Doubao-Seed-2.0-lite",
        "method": "complete_with_tools",
        "purpose": (
            "包装轨自然语言改字幕/标题/贴纸文字 + 调 BGM 音量。"
            "用户描述像『把字幕改成 ××』『BGM 调小一点』。"
            "渲染态不锁(包装可以渲染中改);找不到 item_id 时回落到第一个非 transition 项。"
        ),
        "input": [
            "system: _SYSTEM_PACKAGING",
            "user: 当前 packaging_track (item_id, kind, text) + BGM 当前音量 + 用户指令",
            "tools: _TOOLS_PACKAGING (2 个原子 tool)",
        ],
        "output": (
            "tool_calls: [{\"name\": \"update_packaging_text\"|\"update_bgm_volume\", \"arguments\": {...}}]"
        ),
        "constraints": [
            "update_bgm_volume: volume clamp 到 [0.0, 1.0]",
            "update_packaging_text: 找不到 item_id 时回落到第一个 kind!='transition' 的 item",
            "kind=='transition' 的旧包装项已废弃,不喂给模型",
            "无 tool_call 命中 → 兜底改第一个非 transition 项的 text",
        ],
        "options": [
            "tools 名 ∈ {update_packaging_text, update_bgm_volume}",
            "关键词路由:『字幕/标题/文字/改成/写成』→ update_packaging_text;"
            "『BGM/背景音乐/音量/大声/小声/调到』→ update_bgm_volume",
        ],
        "system_prompt": (
            "你是视频剪辑助手,本次只能修改【包装轨】(packaging_track / BGM)。"
            "可选 tool:改字幕/标题/贴纸文字 / 调 BGM 音量。"
            "禁止改 scene 时长、口播、素材 —— 那些是其他轨道的工具。"
            "『字幕 / 标题 / 文字 / 改成 / 写成』→ update_packaging_text;"
            "『BGM / 背景音乐 / 音量 / 大声 / 小声 / 调到』→ update_bgm_volume。"
        ),
        "user_prompt_template": (
            "当前 packaging_track:\n"
            "- {item_id} kind={kind} text={text!r}\n"
            "- ...\n"
            "BGM 当前音量={volume:.2f}\n"
            "(可选) 用户选中:\n[{track}] {start:.1f}-{end:.1f}s target={target_id}\n"
            "用户指令(轨道=packaging):{instruction}\n"
            "输出 edit_tool_calls。"
        ),
    },
    # ---------------------------------------------------------------------
    {
        "no": 11,
        "title": "自然语言编辑 - 口播轨 (Edit · voice track)",
        "module": "Render / routers/edit.apply_edit (track='voice')",
        "file": "server/app/routers/edit.py",
        "model": "Doubao-Seed-2.0-lite",
        "method": "complete_with_tools",
        "purpose": (
            "口播轨改 main_track[i].narration (TTS 朗读用文字稿)。修改后系统自动重合成 wav 覆盖 voiceover_url。"
            "用户描述像『把开场改得更口语化』『sc-2 改简洁点』。"
            "无 tool_call 命中时,默认拼到首段 narration 前。"
        ),
        "input": [
            "system: _SYSTEM_VOICE",
            "user: 当前 main_track 列表(scene_id + narration) + 用户选区 marks + 用户指令",
            "tools: _TOOLS_VOICE (1 个原子 tool)",
        ],
        "output": "tool_calls: [{\"name\": \"edit_scene_narration\", \"arguments\": {scene_id, narration}}]",
        "constraints": [
            "只能改 narration,不能动时长/素材/字幕/BGM",
            "改后异步触发 TTS 重合成 (synthesize_scene_voice),失败仅日志告警",
            "无 tool_call 时兜底:把指令前 60 字拼到首 scene narration 前",
        ],
        "options": [
            "tools 名 ∈ {edit_scene_narration}",
            "关键词路由:『口播/旁白/念白/朗读/口语化/改得更…』→ edit_scene_narration",
        ],
        "system_prompt": (
            "你是视频剪辑助手,本次只能修改【口播轨】(main_track[i].narration),"
            "也就是 TTS 朗读用的文字稿。修改后系统会自动重新合成 wav。"
            "可选 tool:仅 edit_scene_narration。"
            "禁止改时长、字幕、素材、BGM。"
            "『口播 / 旁白 / 念白 / 朗读 / 口语化 / 改得更…』→ edit_scene_narration。"
        ),
        "user_prompt_template": (
            "当前 Plan main_track:\n"
            "- {scene_id} ({section}) src={source_ref} dur={duration:.1f}s narr={narration!r}\n"
            "- ...\n"
            "(可选) 用户选中:\n[{track}] {start:.1f}-{end:.1f}s target={target_id}\n"
            "用户指令(轨道=voice):{instruction}\n"
            "输出 edit_tool_calls。"
        ),
    },
]


# ----------------------- docx 渲染 ------------------------------

def _set_code_style(doc: Document) -> None:
    """注册 / 调整 'Code' 段落样式 (等宽字体 + 浅灰背景由读者本地 Word 主题决定)。"""
    styles = doc.styles
    if "CodeBlock" not in [s.name for s in styles]:
        style = styles.add_style("CodeBlock", 1)  # WD_STYLE_TYPE.PARAGRAPH = 1
        style.font.name = "Consolas"
        style.font.size = Pt(9)


def _add_code_block(doc: Document, text: str) -> None:
    """把多行字符串写成一段段 'CodeBlock' 段落,保持原始换行。"""
    if not text:
        return
    for line in text.splitlines() or [""]:
        p = doc.add_paragraph(line, style="CodeBlock")
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def _add_kv(doc: Document, key: str, value: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{key}:")
    run.bold = True
    p.add_run(" ")
    p.add_run(value)


def _add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def build_doc() -> None:
    doc = Document()
    _set_code_style(doc)

    # ---- 封面 ----
    title = doc.add_heading("Seecript · LLM 调用节点全集", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("爆款结构迁移引擎 · 大语言模型干预点 / 提示词 / 入出参清单")
    r.italic = True
    r.font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "模型: Doubao-Seed-2.0-lite (多模态)  ·  Provider: 火山方舟 ARK  ·  "
        "整理日期: 2026-06-01"
    ).font.size = Pt(9)

    # ---- 概览 ----
    doc.add_heading("概览", level=1)
    doc.add_paragraph(
        "本文档枚举 Seecript 仓库中 (D:\\Seecript) 所有调用大语言模型的功能节点。"
        "每个节点包含:① 引入 LLM 解决什么问题  ② 入参 / 出参 / 格式约束 / 选项枚举  "
        "③ 完整 system + user prompt。"
    )
    doc.add_paragraph(
        f"共计 {len(NODES)} 个节点。模型统一走 LLMClient 抽象,"
        "默认 provider = doubao_ark (OpenAI 兼容 /chat/completions),"
        "未配 Key 时 fallback 到 MockLLMClient。所有节点失败均有规则兜底,绝不静默失败。"
    )

    # 节点速览表
    doc.add_heading("节点速览", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "节点"
    hdr[2].text = "调用方法"
    hdr[3].text = "所在模块"
    for n in NODES:
        row = table.add_row().cells
        row[0].text = str(n["no"])
        row[1].text = n["title"]
        row[2].text = n["method"]
        row[3].text = n["module"]

    doc.add_page_break()

    # ---- 每个节点展开 ----
    for n in NODES:
        doc.add_heading(f"{n['no']}. {n['title']}", level=1)

        _add_kv(doc, "所在模块", n["module"])
        _add_kv(doc, "源文件", n["file"])
        _add_kv(doc, "模型", n["model"])
        _add_kv(doc, "调用方法", n["method"])

        doc.add_heading("1) 引入 LLM 解决什么问题", level=2)
        doc.add_paragraph(n["purpose"])

        doc.add_heading("2) 入参 / 出参 / 格式约束 / 选项", level=2)

        doc.add_paragraph().add_run("传入参数:").bold = True
        _add_bullets(doc, n["input"])

        doc.add_paragraph().add_run("传出参数 (返回 schema):").bold = True
        _add_code_block(doc, n["output"])

        doc.add_paragraph().add_run("格式化限制:").bold = True
        _add_bullets(doc, n["constraints"])

        doc.add_paragraph().add_run("选项 / 枚举:").bold = True
        _add_bullets(doc, n["options"])

        doc.add_heading("3) 提示词内容", level=2)

        doc.add_paragraph().add_run("System prompt:").bold = True
        _add_code_block(doc, n["system_prompt"])

        doc.add_paragraph().add_run("User prompt 模板:").bold = True
        _add_code_block(doc, n["user_prompt_template"])

        doc.add_paragraph()  # 空行隔离

    # ---- 落盘 ----
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"[ok] wrote {OUT}  ({OUT.stat().st_size:,} bytes, {len(NODES)} nodes)")


if __name__ == "__main__":
    build_doc()
